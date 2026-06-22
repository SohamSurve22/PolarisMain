from ..models import (
    DocumentFormat,
    ExtractionWarning,
    ExtractedDocument,
    ExtractedMetadata,
    PageInfo,
    RawDocument,
    EmptyDocumentError,
)
from ..models.errors import ExtractorError
from ..interfaces import ExtractionResult
from .base import AbstractExtractor


class DOCXExtractor(AbstractExtractor):
    def supported_formats(self) -> list[DocumentFormat]:
        return [DocumentFormat.DOCX]

    def extract(self, document: RawDocument) -> ExtractionResult:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            return ExtractionResult(
                error=ExtractorError(
                    error_code="EXTRACTOR_DEPENDENCY_MISSING",
                    stage="extract",
                    message="python-docx is not installed",
                    context={"document_id": document.id},
                )
            )

        try:
            import io
            doc = DocxDocument(io.BytesIO(document.content))
        except Exception as exc:
            return ExtractionResult(
                error=ExtractorError(
                    error_code="DOCX_PARSE_FAILED",
                    stage="extract",
                    message=f"Failed to open DOCX: {exc}",
                    context={"document_id": document.id, "exception": str(exc)},
                )
            )

        warnings: list[ExtractionWarning] = []
        paragraphs: list[str] = []
        has_tables = False
        has_images = False

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        for table in doc.tables:
            has_tables = True
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                row_text = " | ".join(row_cells)
                if row_text.strip():
                    paragraphs.append(f"[TABLE ROW] {row_text}")

        if doc.inline_shapes:
            has_images = True

        full_text = "\n".join(paragraphs)

        if not full_text.strip():
            return ExtractionResult(
                error=EmptyDocumentError(
                    message="This document appears to be empty",
                    context={"document_id": document.id},
                )
            )

        detected_lang, lang_conf = self._detect_language(full_text)

        metadata = self._compute_metadata(
            full_text=full_text,
            pages=[],
            language=detected_lang,
            language_confidence=lang_conf,
            extraction_strategy="python-docx",
            has_images=has_images,
            has_tables=has_tables,
        )

        extracted = self._build_extracted_document(
            raw_id=document.id,
            format=document.format,
            full_text=full_text,
            pages=[],
            metadata=metadata,
            warnings=warnings,
        )

        return ExtractionResult(document=extracted)
