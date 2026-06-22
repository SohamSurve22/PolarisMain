from __future__ import annotations

import hashlib
from datetime import datetime, timezone
from pathlib import Path

import pytest

from backend.app.services.document_processing.models import (
    DocumentFormat,
    RawDocument,
)


def _make_raw_document(content: bytes, filename: str, format: DocumentFormat) -> RawDocument:
    return RawDocument(
        id=f"test-{hashlib.md5(content).hexdigest()[:12]}",
        filename=filename,
        format=format,
        content=content,
        size_bytes=len(content),
        checksum_sha256=hashlib.sha256(content).hexdigest(),
        upload_timestamp=datetime.now(timezone.utc),
    )


def _read_or_create(directory: str, filename: str, create_fn) -> bytes:
    path = Path(__file__).resolve().parent / "fixtures" / "documents" / directory / filename
    if path.exists():
        return path.read_bytes()
    content = create_fn()
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(content)
    return content


# ─── Fixture files (TXTs are created inline; PDF/DOCX programmatically) ───


@pytest.fixture(scope="session")
def sample_pdf_content() -> bytes:
    def _create() -> bytes:
        import fitz
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text(
            (72, 72),
            "Privacy Policy\n\n"
            "1. Introduction\n"
            "This Privacy Policy describes how we collect, use, and process your personal data.\n\n"
            "2. Data Collection\n"
            "We collect the following categories of personal data:\n"
            "  (a) Name and contact information\n"
            "  (b) Financial information\n"
            "  (c) Usage data\n\n"
            "3. Purpose of Processing\n"
            "Your data is processed for the following purposes:\n"
            "  - Service delivery\n"
            "  - Compliance with legal obligations\n"
            "  - Fraud prevention\n\n"
            "4. Data Retention\n"
            "We retain personal data for as long as necessary to fulfill the purposes described above.\n\n"
            "5. Your Rights\n"
            "Under the Digital Personal Data Protection Act 2023, you have the following rights:\n"
            "  (a) Right to access\n"
            "  (b) Right to correction\n"
            "  (c) Right to erasure\n"
            "  (d) Right to grievance redressal",
            fontsize=11,
        )
        data = doc.tobytes()
        doc.close()
        return data

    return _read_or_create("pdf", "sample_privacy_policy.pdf", _create)


@pytest.fixture(scope="session")
def sample_docx_content() -> bytes:
    def _create() -> bytes:
        from docx import Document
        doc = Document()
        doc.add_heading("Terms of Service", level=1)
        doc.add_paragraph("These Terms of Service govern your use of our platform.")
        doc.add_heading("1. Acceptance of Terms", level=2)
        doc.add_paragraph("By accessing or using our services, you agree to be bound by these Terms.")
        doc.add_heading("2. User Obligations", level=2)
        doc.add_paragraph("You are responsible for maintaining the confidentiality of your account credentials.")
        doc.add_heading("3. Limitation of Liability", level=2)
        doc.add_paragraph("We shall not be liable for any indirect, incidental, or consequential damages.")
        doc.add_heading("4. Governing Law", level=2)
        doc.add_paragraph("These Terms shall be governed by the laws of India.")
        import io
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    return _read_or_create("docx", "sample_terms.docx", _create)


@pytest.fixture(scope="session")
def sample_txt_content() -> bytes:
    def _create() -> bytes:
        return (
            "Data Retention Policy\n"
            "====================\n\n"
            "Effective Date: January 1, 2024\n\n"
            "1. Purpose\n"
            "This Data Retention Policy outlines the periods for which we retain personal data.\n\n"
            "2. Retention Periods\n"
            "  (a) Account data: Retained for the duration of the account plus 3 years\n"
            "  (b) Transaction data: Retained for 7 years as required by law\n"
            "  (c) Log data: Retained for 12 months\n\n"
            "3. Data Deletion\n"
            "Upon expiry of the retention period, data is securely deleted.\n\n"
            "4. Compliance\n"
            "This policy complies with the IT Act 2000 and SPDI Rules 2011.\n"
        ).encode("utf-8")
    return _read_or_create("txt", "sample_retention_policy.txt", _create)


@pytest.fixture(scope="session")
def sample_html_content() -> bytes:
    def _create() -> bytes:
        return (
            "<!DOCTYPE html>\n"
            "<html><head><title>Cookie Policy</title></head>\n"
            "<body>\n"
            "<h1>Cookie Policy</h1>\n"
            "<p>This Cookie Policy explains how we use cookies and similar technologies.</p>\n"
            "<h2>1. What Are Cookies</h2>\n"
            "<p>Cookies are small text files stored on your device when you visit a website.</p>\n"
            "<h2>2. Types of Cookies We Use</h2>\n"
            "<ul>\n"
            "<li>Essential cookies: Required for website functionality</li>\n"
            "<li>Analytics cookies: Help us understand usage patterns</li>\n"
            "<li>Preference cookies: Remember your settings</li>\n"
            "</ul>\n"
            "<h2>3. Managing Cookies</h2>\n"
            "<p>You can control cookies through your browser settings.</p>\n"
            "<h2>4. Contact Us</h2>\n"
            "<p>For questions about this policy, contact dpo@example.com.</p>\n"
            "</body></html>\n"
        ).encode("utf-8")
    return _read_or_create("html", "sample_cookie_policy.html", _create)


@pytest.fixture(scope="session")
def empty_content() -> bytes:
    return b""


@pytest.fixture(scope="session")
def corrupted_pdf_content() -> bytes:
    return b"This is not a PDF file\x00\xFF\xFE\xFD\xFC"


@pytest.fixture(scope="session")
def multi_page_pdf_content() -> bytes:
    def _create() -> bytes:
        import fitz
        doc = fitz.open()
        header_text = "Privacy Policy - CONFIDENTIAL"
        footer_text = "Page {n} | © 2024 Company Name"
        body_text = (
            "1. Introduction\n"
            "This Privacy Policy describes how we collect and process your personal data.\n\n"
            "2. Data We Collect\n"
            "We collect name, email address, phone number, and usage data.\n\n"
            "3. Legal Basis\n"
            "Processing is based on consent, contract performance, and legal obligations.\n\n"
            "4. Data Sharing\n"
            "We share data with service providers and regulatory authorities as required.\n\n"
            "5. International Transfers\n"
            "Data may be transferred to countries with adequate data protection laws."
        )
        for page_num in range(5):
            page = doc.new_page()
            page.insert_text((72, 30), header_text, fontsize=8)
            page.insert_text((72, 72), body_text, fontsize=10)
            page.insert_text((72, 770), footer_text.format(n=page_num + 1), fontsize=8)
        data = doc.tobytes()
        doc.close()
        return data
    return _read_or_create("pdf", "multi_page_policy.pdf", _create)


# ─── RawDocument fixtures ───


@pytest.fixture(scope="session")
def raw_pdf(sample_pdf_content) -> RawDocument:
    return _make_raw_document(sample_pdf_content, "policy.pdf", DocumentFormat.PDF)


@pytest.fixture(scope="session")
def raw_docx(sample_docx_content) -> RawDocument:
    return _make_raw_document(sample_docx_content, "terms.docx", DocumentFormat.DOCX)


@pytest.fixture(scope="session")
def raw_txt(sample_txt_content) -> RawDocument:
    return _make_raw_document(sample_txt_content, "policy.txt", DocumentFormat.TXT)


@pytest.fixture(scope="session")
def raw_html(sample_html_content) -> RawDocument:
    return _make_raw_document(sample_html_content, "policy.html", DocumentFormat.HTML)


@pytest.fixture(scope="session")
def raw_empty() -> RawDocument:
    return _make_raw_document(b"", "empty.pdf", DocumentFormat.PDF)


@pytest.fixture(scope="session")
def raw_corrupted_pdf(corrupted_pdf_content) -> RawDocument:
    return _make_raw_document(corrupted_pdf_content, "corrupted.pdf", DocumentFormat.PDF)
